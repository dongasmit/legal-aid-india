import os
import io
from dotenv import load_dotenv
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser, JsonOutputParser
from langchain_core.runnables import RunnablePassthrough, RunnableParallel
from docx import Document
from fpdf import FPDF
import fitz  # PyMuPDF
from PIL import Image

from typing import Optional
from config import INDEX_NAME, EMBEDDING_MODEL, LLM_MODEL, get_logger

logger = get_logger(__name__)

load_dotenv()

# --- LAZY RESOURCE INITIALISATION ---
# Resources are initialised on first use so that a missing API key does not
# crash the Streamlit app at import time.  Callers should check _init_error
# before invoking any AI function.

_init_error: Optional[str] = None
embeddings = None
vector_db = None
retriever = None
llm = None


def _initialize() -> None:
    """Initialise cloud AI resources.  Stores any error in _init_error instead
    of raising, so the Streamlit app can surface a friendly message."""
    global _init_error, embeddings, vector_db, retriever, llm

    if embeddings is not None or _init_error is not None:
        return  # already attempted

    groq_key = os.getenv("GROQ_API_KEY")
    pinecone_key = os.getenv("PINECONE_API_KEY")

    if not groq_key or not pinecone_key:
        _init_error = (
            "❌ API Keys missing! Check your .env file for "
            "GROQ_API_KEY and PINECONE_API_KEY."
        )
        logger.error(_init_error)
        return

    try:
        embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
        vector_db = PineconeVectorStore(index_name=INDEX_NAME, embedding=embeddings)
        retriever = vector_db.as_retriever(search_kwargs={"k": 5})
        llm = ChatGroq(temperature=0.1, model_name=LLM_MODEL, api_key=groq_key)
        logger.info("AI resources initialised successfully.")
    except Exception as exc:
        _init_error = f"❌ Failed to initialise AI resources: {exc}"
        logger.error(_init_error)


# --- HELPER FUNCTIONS ---
def generate_docx(text):
    doc = Document()
    doc.add_heading('JurisOne Legal Draft', 0)
    for paragraph in text.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_pdf(text):
    """Generate a PDF from *text* with proper UTF-8 support via fpdf2."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    pdf.multi_cell(0, 8, text)
    buffer = io.BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer


def get_source_image(file_path, page_number):
    try:
        # Normalize the path so it works perfectly on both Windows (\) and Linux (/)
        local_path = os.path.normpath(file_path)

        # If the file doesn't exist (like when running on Streamlit Cloud), trigger text fallback
        if not os.path.exists(local_path):
            return None

        doc = fitz.open(local_path)

        # Ensure page number is an integer to avoid indexing errors
        page_idx = int(page_number)
        if page_idx >= len(doc):
            return None

        page = doc.load_page(page_idx)
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        return img
    except Exception as exc:
        logger.warning("Image generation skipped: %s", exc)
        return None


# --- INTELLIGENCE FUNCTIONS ---

# 1. THE STRATEGIST (Research with Memory)
def get_research_response(query, history_text):
    _initialize()
    if _init_error:
        raise RuntimeError(_init_error)

    query_transform_prompt = ChatPromptTemplate.from_template(
        """
        Given the conversation history and the new question, create a precise search query.

        HISTORY: {history}
        NEW QUESTION: {question}

        OUTPUT ONLY THE SEARCH QUERY.
        """
    )
    search_query_chain = query_transform_prompt | llm | StrOutputParser()
    generated_query = search_query_chain.invoke({"history": history_text, "question": query})

    answer_prompt = ChatPromptTemplate.from_template(
        """
        You are a Senior Legal Partner. Provide strategic advice.

        CONTEXT (Laws/Judgments): {context}
        USER QUERY: {question}

        **INSTRUCTIONS:**
        1. Answer based on the CONTEXT provided.
        2. If the user asks for documents, list them clearly.
        3. END your response by saying:
           "I can draft these for you. Just say: 'Draft the [Document Name]'."

        **FORMAT:**
        - **Executive Summary**
        - **Legal Provisions** (Cite Sections)
        - **Precedents** (Cite Case Names if in context)
        - **Strategic Steps**
        """
    )

    rag_chain = (
        RunnableParallel({
            "context": lambda x: retriever.invoke(generated_query),
            "question": RunnablePassthrough()
        })
        .assign(answer=answer_prompt | llm | StrOutputParser())
    )

    return rag_chain.invoke(query)


# 2. THE INTERVIEWER (Context Aware)
def analyze_drafting_needs(user_input, history_text):
    _initialize()
    if _init_error:
        raise RuntimeError(_init_error)

    analyzer_prompt = ChatPromptTemplate.from_template(
        """
        You are a Legal Drafting Expert.
        Analyze the HISTORY to understand what case we are dealing with.

        Current Request: {input}
        Full Conversation History: {history}

        **TASK:**
        1. Extract the Case Type from history.
        2. Identify the document the user wants to draft now.
        3. Check if we have names/dates/details.

        **OUTPUT JSON ONLY:**
        {{
            "status": "READY" or "MISSING_INFO",
            "missing_details": ["List of questions"],
            "document_type": "Specific Document Name"
        }}
        """
    )
    chain = analyzer_prompt | llm | JsonOutputParser()
    return chain.invoke({"input": user_input, "history": history_text})


# 3. THE DRAFTER
def generate_legal_draft(user_input, history_text, doc_type):
    _initialize()
    if _init_error:
        raise RuntimeError(_init_error)

    draft_prompt = ChatPromptTemplate.from_template(
        """
        You are a Senior Advocate. Draft a professional **{doc_type}**.
        Use the details from this history:
        {history}

        Requirements:
        - Full Legal Format.
        - Use placeholders [_______] for missing info.
        - NO conversational text. Just the document content.
        """
    )
    chain = draft_prompt | llm | StrOutputParser()
    return chain.invoke({"history": history_text, "input": user_input, "doc_type": doc_type})


def convert_law_code(query):
    _initialize()
    if _init_error:
        raise RuntimeError(_init_error)

    converter_prompt = ChatPromptTemplate.from_template(
        "Map Old IPC '{query}' to New BNS. Output: Old -> New (Key Changes)."
    )
    chain = converter_prompt | llm | StrOutputParser()
    return chain.invoke({"query": query})


# --- MAIN ROUTER ---
def ask_legal_ai(user_input, chat_history_list):
    _initialize()
    if _init_error:
        raise RuntimeError(_init_error)

    history_text = "\n".join([f"{msg['role'].upper()}: {msg['content']}" for msg in chat_history_list])

    draft_keywords = ["draft", "write", "prepare", "create", "generate"]
    is_draft_request = any(k in user_input.lower() for k in draft_keywords)

    if is_draft_request:
        analysis = analyze_drafting_needs(user_input, history_text)

        if analysis["status"] == "MISSING_INFO":
            questions = "\n".join([f"- {q}" for q in analysis["missing_details"]])
            return {
                "type": "interview",
                "answer": f"**Drafting Protocol: {analysis['document_type']}**\n\nI have the legal context, but I need specific details to fill the document:\n\n{questions}",
                "context": []
            }
        else:
            draft_text = generate_legal_draft(user_input, history_text, analysis["document_type"])
            return {
                "type": "draft",
                "answer": f"**Draft Ready: {analysis['document_type']}**\n\nHere is the legally compliant draft based on our case strategy.",
                "docx": generate_docx(draft_text),
                "pdf": generate_pdf(draft_text),
                "context": []
            }

    else:
        response = get_research_response(user_input, history_text)
        return {
            "type": "research",
            "answer": response["answer"],
            "context": response["context"]
        }